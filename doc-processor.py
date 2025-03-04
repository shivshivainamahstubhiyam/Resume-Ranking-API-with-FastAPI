import os
import tempfile
from typing import Union

import docx
import PyPDF2
from fastapi import UploadFile


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from PDF or DOCX files.
    
    Args:
        file: The uploaded file object
        
    Returns:
        str: Extracted text from the document
    
    Raises:
        ValueError: If file format is not supported
    """
    content = await file.read()
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # Extract text based on file type
        if file_ext == ".pdf":
            text = extract_text_from_pdf(temp_file_path)
        elif file_ext == ".docx":
            text = extract_text_from_docx(temp_file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        return text
    finally:
        # Clean up the temporary file
        os.unlink(temp_file_path)
        # Reset file pointer for potential reuse
        await file.seek(0)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        str: Extracted text
    """
    text = ""
    with open(file_path, "rb") as f:
        pdf_reader = PyPDF2.PdfReader(f)
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
    
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text
    """
    doc = docx.Document(file_path)
    text = ""
    
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text
